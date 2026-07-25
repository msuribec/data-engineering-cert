#!/usr/bin/env python3
"""Parse AWS Data Engineer study materials into a single study_data.json.

Sources:
  - Flashcards_QA.csv           -> Q/A flashcards
  - Flashcards_AWS_Services.csv -> service flashcards
  - Study Materials/**/*Quiz*.docx -> quizzes (MC + short answer + answer key)
  - Matching study sheets       -> explicit primary DEA-C01 domain per quiz
"""
import csv, glob, json, os, re
from collections.abc import Callable
from typing import Any

from docx import Document

BASE = "Study Materials"
FLASHCARDS_DIR = os.path.join(BASE, "Flashcards")
QUIZ_DIR = os.path.join(BASE, "Study sheets and Quiz")
OUT = os.path.join( "data", "study_data.json")
DOMAIN_RE = re.compile(r"Primary domain:\s*([1-4])", re.I)


def load_csv(path: str) -> list[dict[str, str]]:
    rows = []
    if not os.path.exists(path):
        return rows
    with open(path, newline="", encoding="utf-8-sig") as f:
        for r in csv.DictReader(f):
            front = (r.get("Front") or "").strip()
            back = (r.get("Back") or "").strip()
            tags = (r.get("Tags") or "").strip()
            if not front:
                continue
            section = ""
            m = re.search(r"Section-([0-9]+)-([A-Za-z0-9-]+)", tags)
            if m:
                section = f"Section {m.group(1)} - " + m.group(2).replace("-", " ")
            rows.append({"front": front, "back": back, "tags": tags, "section": section})
    return rows


OPT_RE = re.compile(r"^([A-D])\)\s*(.*)$")
Q_RE = re.compile(r"^Q(\d+)\.\s*(.*)$")
SA_RE = re.compile(r"^(\d+)\.\s*(.*)$")
CORRECT_RE = re.compile(r"^Q(\d+)\s*[—-]\s*Correct:\s*([A-D])", re.I)


def parse_quiz(path: str) -> dict[str, Any]:
    d = Document(path)
    lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    title = lines[0] if lines else os.path.basename(path)

    # Segment into regions
    def idx(pred: Callable[[str], bool], start: int = 0) -> int:
        for i in range(start, len(lines)):
            if pred(lines[i]):
                return i
        return -1

    i_mc = idx(lambda t: t.startswith("Part A"))
    i_sa = idx(lambda t: t.startswith("Part B"))
    i_key = idx(lambda t: t.startswith("Answer Key"))

    mc = {}       # num -> {question, options{}}
    order = []
    # --- MC questions (Part A region) ---
    region = lines[i_mc:i_sa] if i_mc >= 0 else []
    cur = None
    for t in region:
        qm = Q_RE.match(t)
        if qm:
            cur = int(qm.group(1))
            mc[cur] = {"question": qm.group(2), "options": {}}
            order.append(cur)
            continue
        om = OPT_RE.match(t)
        if om and cur is not None:
            mc[cur]["options"][om.group(1)] = om.group(2).strip()

    # --- Short answer questions (Part B region) ---
    sa = {}
    region = lines[i_sa:i_key] if i_sa >= 0 else []
    for t in region:
        if t.startswith("Part B") or t.lower().startswith("scenario"):
            continue
        sm = SA_RE.match(t)
        if sm:
            sa[int(sm.group(1))] = {"question": sm.group(2), "answer": ""}

    # --- Answer key region ---
    key = lines[i_key:] if i_key >= 0 else []
    j = 0
    mode = "mc"
    while j < len(key):
        t = key[j]
        if t.startswith("Short Answer"):
            mode = "sa"
        cm = CORRECT_RE.match(t)
        if cm and mode == "mc":
            num = int(cm.group(1))
            expl_parts = []
            j += 1
            while j < len(key) and not CORRECT_RE.match(key[j]) and not key[j].startswith("Short Answer"):
                kt = key[j]
                if kt.startswith("Why the correct answer is right:"):
                    expl_parts.append(kt.split(":", 1)[1].strip())
                j += 1
            if num in mc:
                mc[num]["correct"] = cm.group(2).upper()
                mc[num]["explanation"] = " ".join(expl_parts)
            continue
        # short answer key: "1.  answer text"
        if mode == "sa":
            sm = SA_RE.match(t)
            if sm:
                n = int(sm.group(1))
                if n in sa:
                    sa[n]["answer"] = sm.group(2)
        j += 1

    questions = []
    for n in order:
        q = mc[n]
        if "correct" in q and len(q["options"]) >= 2:
            questions.append({
                "type": "mc",
                "question": q["question"],
                "options": q["options"],
                "correct": q["correct"],
                "explanation": q.get("explanation", ""),
            })
    short = [{"type": "short", "question": v["question"], "answer": v["answer"]}
             for k, v in sorted(sa.items()) if v["question"]]
    return {"title": title, "mc": questions, "short": short}


def study_sheet_domain(quiz_path: str) -> int:
    """Read the explicit primary DEA-C01 domain from the matching study sheet."""
    sheet_path = re.sub(r"\s*-\s*Quiz\.docx$", " - Study Sheet.docx", quiz_path)
    if not os.path.exists(sheet_path):
        raise ValueError(f"matching study sheet not found: {sheet_path}")
    document = Document(sheet_path)
    for paragraph in document.paragraphs:
        match = DOMAIN_RE.search(paragraph.text)
        if match:
            return int(match.group(1))
    raise ValueError(f"primary domain mapping not found in {sheet_path}")


def main() -> None:
    qa = load_csv(os.path.join(FLASHCARDS_DIR, "Flashcards_QA.csv"))
    svc = load_csv(os.path.join(FLASHCARDS_DIR, "Flashcards_AWS_Services.csv"))
    terms = load_csv(os.path.join(FLASHCARDS_DIR, "Flashcards_Terms.csv"))

    quizzes = []
    files = sorted(glob.glob(os.path.join(QUIZ_DIR, "*", "*Quiz*.docx")))
    for f in files:
        section = os.path.basename(os.path.dirname(f))
        topic = re.sub(r"\s*-\s*Quiz$", "", os.path.splitext(os.path.basename(f))[0])
        topic = re.sub(r"^\d+\s*-\s*", "", topic)
        try:
            parsed = parse_quiz(f)
            domain = study_sheet_domain(f)
        except Exception as e:
            print(f"  ! failed {f}: {e}")
            continue
        if parsed["mc"] or parsed["short"]:
            quizzes.append({
                "section": section,
                "topic": topic,
                "title": parsed["title"],
                "domain": domain,
                "mc": parsed["mc"],
                "short": parsed["short"],
            })

    sections = sorted({q["section"] for q in quizzes},
                      key=lambda s: int(re.search(r"\d+", s).group()) if re.search(r"\d+", s) else 0)

    data = {
        "flashcards_qa": qa,
        "flashcards_services": svc,
        "flashcards_terms": terms,
        "quizzes": quizzes,
        "sections": sections,
    }
    with open(OUT, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)

    total_mc = sum(len(q["mc"]) for q in quizzes)
    total_sa = sum(len(q["short"]) for q in quizzes)
    print(f"Q/A flashcards:      {len(qa)}")
    print(f"Service flashcards:  {len(svc)}")
    print(f"Term flashcards:     {len(terms)}")
    print(f"Quizzes:             {len(quizzes)} across {len(sections)} sections")
    print(f"  MC questions:      {total_mc}")
    print(f"  Short answers:     {total_sa}")
    print(f"Wrote {OUT} ({os.path.getsize(OUT)//1024} KB)")


if __name__ == "__main__":
    main()
